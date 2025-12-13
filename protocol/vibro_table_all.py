import pandas as pd
import re
from .utils.utils_ppu_testus import vibraTableOne
import zipfile
import rarfile
import tempfile
import os
import shutil
from pprint import pprint
import io
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

import matplotlib.pyplot as plt
import shutil

import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from matplotlib import rcParams


plt.style.use('seaborn-v0_8-whitegrid')
rcParams['figure.facecolor'] = 'white'
rcParams['savefig.facecolor'] = 'white'
rcParams['font.family'] = 'serif'
rcParams['font.sans-serif'] = ['Times New Roman']
rcParams['font.style'] = 'normal'
plt.rcParams['mathtext.default'] = 'regular'


def create_comparison_plots(samples_data, doc):
    """
    Создает сравнительные графики разных образцов при одинаковых пригрузах
    """
    # Собираем все доступные массы пригрузов
    all_masses = set()
    for sample_id, data in samples_data.items():
        if 'results' in data:
            all_masses.update([float(mass) for mass in data['results'].keys()])
    
    all_masses = sorted(all_masses)
    
    if not all_masses:
        return
    
    # 1. Сравнение амплитудно-частотных характеристик (модуль передаточной функции)
    doc.add_page_break()
    heading = doc.add_heading('Сравнение амплитудно-частотных характеристик образцов', level=1)
    apply_times_new_roman_12_black(heading)
    
    # Создаем графики для каждой массы пригруза
    for mass in all_masses:
        mass_str = str(mass)
        
        # Проверяем, есть ли данные для этой массы у всех образцов
        has_data_for_all = True
        for sample_id, data in samples_data.items():
            if 'datas' not in data or mass_str not in data['datas']:
                has_data_for_all = False
                break
        
        if not has_data_for_all:
            continue
        
        # Создаем график сравнения модуля передаточной функции
        fig, ax = plt.subplots(figsize=(8, 5), tight_layout=True)
        
        # Настраиваем цвета для разных образцов
        colors = plt.cm.tab10(np.linspace(0, 1, len(samples_data)))
        
        for i, (sample_id, data) in enumerate(samples_data.items()):
            if 'datas' in data and mass_str in data['datas']:
                freq, tf_module, _ = data['datas'][mass_str]
                
                # Рисуем график для текущего образца
                ax.plot(freq, tf_module, 
                       linewidth=2, 
                       color=colors[i],
                       label=f'Образец {sample_id}')
        
        # Настройка графика
        ax.set_xlabel('Частота, Гц', fontsize=14, fontname='Times New Roman', color='black')
        ax.set_ylabel('Модуль передаточной функции', fontsize=14, fontname='Times New Roman', color='black')
        
        # Добавляем сетку и легенду
        ax.grid(visible=True, which='both', axis='both', ls='--')
        ax.legend(loc='lower right', fontsize=10, frameon=True)
        
        # Устанавливаем шрифт Times New Roman черный для всех текстовых элементов
        for item in ([ax.title, ax.xaxis.label, ax.yaxis.label] + 
                     ax.get_xticklabels() + ax.get_yticklabels()):
            item.set_fontname('Times New Roman')
            item.set_fontsize(10)
            item.set_color('black')
        
        plt.tight_layout()
        
        # Сохраняем график в буфер
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        buf.seek(0)
        
        # Добавляем график в документ
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(buf, width=Inches(6))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption = doc.add_paragraph(f"Рисунок: Сравнение АЧХ образцов при пригрузе {mass} кг", style='Caption')
        apply_times_new_roman_12_black_to_paragraph(caption)

        buf.close()
        
        # Закрываем figure чтобы освободить память
        plt.close(fig)
    
    # 2. Сравнение графиков эффективности виброизоляции
    doc.add_page_break()
    heading = doc.add_heading('Сравнение эффективности виброизоляции образцов', level=1)
    apply_times_new_roman_12_black(heading)
    
    for mass in all_masses:
        mass_str = str(mass)
        
        # Проверяем, есть ли данные для этой массы у всех образцов
        has_data_for_all = True
        for sample_id, data in samples_data.items():
            if 'datas' not in data or mass_str not in data['datas']:
                has_data_for_all = False
                break
        
        if not has_data_for_all:
            continue
        
        # Создаем график сравнения эффективности виброизоляции
        fig, ax = plt.subplots(figsize=(8, 5), tight_layout=True)
        
        # Настраиваем цвета для разных образцов
        colors = plt.cm.tab10(np.linspace(0, 1, len(samples_data)))
        
        for i, (sample_id, data) in enumerate(samples_data.items()):
            if 'datas' in data and mass_str in data['datas']:
                freq, _, isolation_eff = data['datas'][mass_str]
                
                # Рисуем график для текущего образца
                ax.plot(freq, isolation_eff, 
                       linewidth=2, 
                       color=colors[i],
                       label=f'Образец {sample_id}')
        
        # Настройка графика
        ax.set_xlabel('Частота, Гц', fontsize=14, fontname='Times New Roman', color='black')
        ax.set_ylabel('Эффективность виброизоляции, дБ', fontsize=14, fontname='Times New Roman', color='black')
        
        # Добавляем сетку и легенду
        ax.grid(visible=True, which='both', axis='both', ls='--')
        ax.legend(loc='lower right', fontsize=10, frameon=True)
        
        # Устанавливаем шрифт Times New Roman для всех текстовых элементов
        for item in ([ax.title, ax.xaxis.label, ax.yaxis.label] + 
                     ax.get_xticklabels() + ax.get_yticklabels()):
            item.set_fontname('Times New Roman')
            item.set_fontsize(10)
            item.set_color('black')
        
        plt.tight_layout()
        
        # Сохраняем график в буфер
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        buf.seek(0)
        
        # Добавляем график в документ
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(buf, width=Inches(6))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption = doc.add_paragraph(f"Рисунок: Сравнение эффективности виброизоляции при пригрузе {mass} кг", style='Caption')
        apply_times_new_roman_12_black_to_paragraph(caption)

        buf.close()
        
        # Закрываем figure чтобы освободить память
        plt.close(fig)
    
    # 3. Сравнение резонансных частот для разных пригрузов
    doc.add_page_break()
    heading = doc.add_heading('Сравнение резонансных частот образцов', level=1)
    apply_times_new_roman_12_black(heading)
    
    # Создаем график зависимости резонансной частоты от массы пригруза
    fig, ax = plt.subplots(figsize=(10, 6), tight_layout=True)
    
    # Настраиваем цвета для разных образцов
    colors = plt.cm.tab10(np.linspace(0, 1, len(samples_data)))
    
    for i, (sample_id, data) in enumerate(samples_data.items()):
        if 'results' not in data:
            continue
        
        # Собираем данные для графика
        masses = []
        frequencies = []
        
        for mass_str, values in data['results'].items():
            pressure, Fpeak, Ed, damp = values
            masses.append(float(mass_str))
            frequencies.append(float(Fpeak))
        
        # Сортируем по массе
        sorted_data = sorted(zip(masses, frequencies))
        if sorted_data:
            masses_sorted, freqs_sorted = zip(*sorted_data)
            
            # Рисуем график для текущего образца
            ax.plot(masses_sorted, freqs_sorted, 
                   marker='o', 
                   linestyle='-', 
                   linewidth=2,
                   markersize=6,
                   color=colors[i],
                   label=f'Образец {sample_id}')
    
    # Настройка графика
    ax.set_xlabel('Масса пригруза, кг', fontsize=12, fontname='Times New Roman', color='black')
    ax.set_ylabel('Резонансная частота, Гц', fontsize=12, fontname='Times New Roman', color='black')
    ax.set_title('Зависимость резонансной частоты от массы пригруза', 
                fontsize=14, fontname='Times New Roman', fontweight='bold', color='black')
    
    # Добавляем сетку и легенду
    ax.grid(visible=True, which='both', axis='both', ls='--')
    ax.legend(loc='lower right', fontsize=10, frameon=True)
    
    # Устанавливаем шрифт Times New Roman для всех текстовых элементов
    for item in ([ax.title, ax.xaxis.label, ax.yaxis.label] + 
                 ax.get_xticklabels() + ax.get_yticklabels()):
        item.set_fontname('Times New Roman')
        item.set_fontsize(10)
        item.set_color('black')
    
    plt.tight_layout()
    
    # Сохраняем график в буфер
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    buf.seek(0)
    
    # Добавляем график в документ
    caption = doc.add_paragraph("Рисунок: Зависимость резонансной частоты от массы пригруза", style='Caption')
    apply_times_new_roman_12_black_to_paragraph(caption)
    
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(buf, width=Inches(6))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buf.close()
    
    # Закрываем figure чтобы освободить память
    plt.close(fig)
    
    # 4. Сравнение динамического модуля упругости для разных пригрузов
    doc.add_page_break()
    heading = doc.add_heading('Сравнение динамического модуля упругости образцов', level=1)
    apply_times_new_roman_12_black(heading)
    
    # Создаем график зависимости динамического модуля упругости от массы пригруза
    fig, ax = plt.subplots(figsize=(10, 6), tight_layout=True)
    
    # Настраиваем цвета для разных образцов
    colors = plt.cm.tab10(np.linspace(0, 1, len(samples_data)))
    
    for i, (sample_id, data) in enumerate(samples_data.items()):
        if 'results' not in data:
            continue
        
        # Собираем данные для графика
        masses = []
        ed_values = []
        
        for mass_str, values in data['results'].items():
            pressure, Fpeak, Ed, damp = values
            masses.append(float(mass_str))
            ed_values.append(float(Ed))
        
        # Сортируем по массе
        sorted_data = sorted(zip(masses, ed_values))
        if sorted_data:
            masses_sorted, ed_sorted = zip(*sorted_data)
            
            # Рисуем график для текущего образца
            ax.plot(masses_sorted, ed_sorted, 
                   marker='s', 
                   linestyle='-', 
                   linewidth=2,
                   markersize=6,
                   color=colors[i],
                   label=f'Образец {sample_id}')
    
    # Настройка графика
    ax.set_xlabel('Масса пригруза, кг', fontsize=12, fontname='Times New Roman', color='black')
    ax.set_ylabel('Динамический модуль упругости, Н/мм²', fontsize=12, fontname='Times New Roman', color='black')
    ax.set_title('Зависимость динамического модуля упругости от массы пригруза', 
                fontsize=14, fontname='Times New Roman', fontweight='bold', color='black')
    
    # Добавляем сетку и легенду
    ax.grid(visible=True, which='both', axis='both', ls='--')
    ax.legend(loc='lower right', fontsize=10, frameon=True)
    
    # Устанавливаем шрифт Times New Roman для всех текстовых элементов
    for item in ([ax.title, ax.xaxis.label, ax.yaxis.label] + 
                 ax.get_xticklabels() + ax.get_yticklabels()):
        item.set_fontname('Times New Roman')
        item.set_fontsize(10)
        item.set_color('black')
    
    plt.tight_layout()
    
    # Сохраняем график в буфер
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    buf.seek(0)
    
    # Добавляем график в документ
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(buf, width=Inches(6))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph("Рисунок: Зависимость динамического модуля упругости от массы пригруза", style='Caption')
    apply_times_new_roman_12_black_to_paragraph(caption)

    buf.close()
    
    # Закрываем figure чтобы освободить память
    plt.close(fig)

def create_full_report(samples_data, output_file='Оранжевый.docx'):
    """
    Создает полный отчет с таблицами и картинками
    """
    doc = Document()
    
    # Установка стиля Times New Roman 12pt черный для всего документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = None  # Черный цвет
    
    # 1. Геометрические характеристики (таблица)
    heading = doc.add_heading('Геометрические характеристики образцов', level=1)
    apply_times_new_roman_12_black(heading)
    
    table1 = doc.add_table(rows=1, cols=5)
    table1.style = 'Table Grid'
    
    headers = ['Образец', 'Высота (мм)', 'Ширина (мм)', 'Длина (мм)', 'Базовая масса (г)']
    for i, header in enumerate(headers):
        cell = table1.cell(0, i)
        cell.text = header
        apply_times_new_roman_black_to_cell(cell)
    
    row_idx = 1
    for sample_id, data in samples_data.items():
        table1.add_row()
        table1.cell(row_idx, 0).text = str(sample_id)
        table1.cell(row_idx, 1).text = str(data['geometric_params']['height'])
        table1.cell(row_idx, 2).text = str(data['geometric_params']['width'])
        table1.cell(row_idx, 3).text = str(data['geometric_params']['length'])
        table1.cell(row_idx, 4).text = str(data['geometric_params']['base_mass'])
        
        # Установка шрифта для данных в таблице
        for col in range(5):
            cell = table1.cell(row_idx, col)
            apply_times_new_roman_black_to_cell(cell)
        
        row_idx += 1
    
    doc.add_page_break()
    
    # 2. Результаты испытаний в формате: Характеристика | Величина пригруза
    heading = doc.add_heading('Результаты испытаний всех образцов', level=1)
    apply_times_new_roman_12_black(heading)
    
    # Собираем все массы для создания общей структуры таблиц
    all_masses = set()
    for sample_id, data in samples_data.items():
        if 'results' in data:
            all_masses.update([float(mass) for mass in data['results'].keys()])
    
    all_masses = sorted(all_masses)
    
    for sample_id, data in samples_data.items():
        if 'results' not in data:
            continue
            
        heading = doc.add_heading(f'Образец {sample_id}', level=2)
        apply_times_new_roman_12_black(heading)
        
        # Создаем таблицу в требуемом формате
        masses = sorted([float(mass) for mass in data['results'].keys()])
        
        # Создаем таблицу: 1 строка заголовков + 4 строки данных
        table2 = doc.add_table(rows=5, cols=len(masses) + 1)
        table2.style = 'Table Grid'
        
        # Заголовок таблицы
        table2.cell(0, 0).text = "Характеристика"
        apply_times_new_roman_black_to_cell(table2.cell(0, 0))
        
        for i, mass in enumerate(masses, 1):
            table2.cell(0, i).text = f"{mass} кг"
            apply_times_new_roman_black_to_cell(table2.cell(0, i))
        
        # Данные для таблицы
        characteristics = [
            "Удельное давление, кПа",
            "Динамический модуль упругости, Н/мм²",
            "Коэффициент потерь", 
            "Частота резонанса, Гц"
        ]
        
        # Заполняем данные
        for row, char_name in enumerate(characteristics, 1):
            table2.cell(row, 0).text = char_name
            apply_times_new_roman_black_to_cell(table2.cell(row, 0))
            
            for col, mass in enumerate(masses, 1):
                pressure, Fpeak, Ed, damp = data['results'][str(mass)]
                if row == 1:  # Удельное давление, кПа
                    value = f"{float(pressure):.2f}"
                elif row == 2:  # Динамический модуль упругости, Н/мм²
                    value = f"{Ed:.2f}"
                elif row == 3:  # Коэффициент потерь
                    value = f"{float(damp):.2f}"
                elif row == 4:  # Частота резонанса, Гц
                    value = f"{float(Fpeak):.2f}"
                
                table2.cell(row, col).text = value
                apply_times_new_roman_black_to_cell(table2.cell(row, col))
        
        doc.add_paragraph()  # Добавляем отступ между таблицами
    
    # 3. Таблица средних значений с коэффициентом вариации
    doc.add_page_break()
    heading = doc.add_heading('Средние значения характеристик по всем образцам', level=1)
    apply_times_new_roman_12_black(heading)
    
    # Определяем, для каких характеристик нужен коэффициент вариации
    characteristics_with_cv = [
        "Динамический модуль упругости, Н/мм²",
        "Коэффициент потерь"
    ]
    
    # Количество строк: 4 основные характеристики + 2 строки для коэффициентов вариации
    total_rows = 4 + len(characteristics_with_cv)
    table_avg = doc.add_table(rows=total_rows, cols=len(all_masses) + 1)
    table_avg.style = 'Table Grid'
    
    # Заголовок таблицы
    table_avg.cell(0, 0).text = "Характеристика"
    apply_times_new_roman_black_to_cell(table_avg.cell(0, 0))
    
    for i, mass in enumerate(all_masses, 1):
        table_avg.cell(0, i).text = f"{mass} кг"
        apply_times_new_roman_black_to_cell(table_avg.cell(0, i))
    
    # Подготовка данных для расчета средних и коэффициента вариации
    characteristics_data = {
        "Удельное давление, кПа": {mass: [] for mass in all_masses},
        "Динамический модуль упругости, Н/мм²": {mass: [] for mass in all_masses},
        "Коэффициент потерь": {mass: [] for mass in all_masses},
        "Частота резонанса, Гц": {mass: [] for mass in all_masses}
    }
    
    # Собираем данные от всех образцов
    for sample_id, data in samples_data.items():
        if 'results' not in data:
            continue
            
        for mass_str, values in data['results'].items():
            mass = float(mass_str)
            if mass in all_masses:
                pressure, Fpeak, Ed, damp = values
                characteristics_data["Удельное давление, кПа"][mass].append(float(pressure))
                characteristics_data["Динамический модуль упругости, Н/мм²"][mass].append(Ed)
                characteristics_data["Коэффициент потерь"][mass].append(float(damp))
                characteristics_data["Частота резонанса, Гц"][mass].append(float(Fpeak))
    
    # Функция для расчета коэффициента вариации
    def calculate_coefficient_of_variation(values):
        if not values or len(values) < 2:
            return None
        mean = sum(values) / len(values)
        std_dev = (sum((x - mean) ** 2 for x in values) / len(values)) ** 0.5
        cv = (std_dev / mean) * 100 if mean != 0 else 0
        return cv
    
    # Заполняем таблицу средними значениями
    characteristics = [
        "Удельное давление, кПа",
        "Динамический модуль упругости, Н/мм²",
        "Коэффициент потерь", 
        "Частота резонанса, Гц"
    ]
    
    current_row = 1
    
    for char_name in characteristics:
        # Проверяем, что текущая строка существует в таблице
        if current_row >= total_rows:
            break
            
        # Средние значения
        table_avg.cell(current_row, 0).text = char_name
        apply_times_new_roman_black_to_cell(table_avg.cell(current_row, 0))
        
        for col, mass in enumerate(all_masses, 1):
            values = characteristics_data[char_name][mass]
            if values:
                avg_value = sum(values) / len(values)
                if char_name == "Удельное давление, кПа":
                    table_avg.cell(current_row, col).text = f"{avg_value:.2f}"
                elif char_name == "Динамический модуль упругости, Н/мм²":
                    table_avg.cell(current_row, col).text = f"{avg_value:.2f}"
                elif char_name == "Коэффициент потерь":
                    table_avg.cell(current_row, col).text = f"{avg_value:.2f}"
                elif char_name == "Частота резонанса, Гц":
                    table_avg.cell(current_row, col).text = f"{avg_value:.2f}"
            else:
                table_avg.cell(current_row, col).text = "Н/Д"
            
            apply_times_new_roman_black_to_cell(table_avg.cell(current_row, col))
        
        # Если это характеристика, для которой нужен коэффициент вариации
        if char_name in characteristics_with_cv:
            current_row += 1
            # Проверяем, что строка для коэффициента вариации существует
            if current_row >= total_rows:
                break
                
            # Коэффициент вариации
            table_avg.cell(current_row, 0).text = "Коэффициент вариации, %"
            apply_times_new_roman_black_to_cell(table_avg.cell(current_row, 0))
            
            for col, mass in enumerate(all_masses, 1):
                values = characteristics_data[char_name][mass]
                if values and len(values) > 1:
                    cv = calculate_coefficient_of_variation(values)
                    table_avg.cell(current_row, col).text = f"{cv:.2f}" if cv is not None else "Н/Д"
                else:
                    table_avg.cell(current_row, col).text = "Н/Д"
                
                apply_times_new_roman_black_to_cell(table_avg.cell(current_row, col))
        
        current_row += 1
    
    # 4. График зависимости динамического модуля упругости от удельного давления (СРЕДНИЕ ЗНАЧЕНИЯ)
    doc.add_page_break()
    heading = doc.add_heading('График зависимости динамического модуля упругости от удельного давления (средние значения)', level=1)
    apply_times_new_roman_12_black(heading)
    
    # Создаем график средних значений
    fig, ax = plt.subplots(figsize=(7, 4), tight_layout=True)
    
    # Собираем средние значения для графика
    avg_pressures_all = []
    avg_ed_all = []
    
    for mass in all_masses:
        pressures = characteristics_data["Удельное давление, кПа"][mass]
        ed_values = characteristics_data["Динамический модуль упругости, Н/мм²"][mass]
        
        if pressures and ed_values:
            avg_pressure = sum(pressures) / len(pressures)
            avg_ed = sum(ed_values) / len(ed_values)
            avg_pressures_all.append(avg_pressure)
            avg_ed_all.append(avg_ed)
    
    # Сортируем данные по давлению для правильного отображения линии
    if avg_pressures_all and avg_ed_all:
        sorted_data = sorted(zip(avg_pressures_all, avg_ed_all))
        pressures_sorted, ed_sorted = zip(*sorted_data)
        
        line = ax.plot(pressures_sorted, ed_sorted, 
                    marker='o', 
                    linestyle='-', 
                    linewidth=1.5,
                    markersize=4,
                    markerfacecolor='blue',
                    markeredgecolor='blue',
                    color='orange',
                    label='Средние значения')
        
        # Добавляем подписи с ограничением области отображения
        for i, (x, y) in enumerate(zip(pressures_sorted, ed_sorted)):
            ax.annotate(f'{y:.2f}', 
                    xy=(x, y), 
                    xytext=(-5, 10),
                    textcoords='offset points',
                    fontsize=8,
                    fontname='Times New Roman',
                    color='black',
                    clip_on=True,  # Важно: не выходить за границы графика
                    )

    # Устанавливаем отступы для графика
    ax.margins(x=0.1, y=0.1)  # Добавляем 10% отступ по краям
    
    # Настройка графика
    ax.set_xlabel('Удельное давление, кПа', fontsize=12, fontname='Times New Roman', color='black')
    ax.set_ylabel('Динамический модуль упругости, Н/мм²', fontsize=12, fontname='Times New Roman', color='black')
    
    # Добавляем сетку
    ax.grid(visible=True, which='both', axis='both', ls='--')
    
    # Устанавливаем шрифт Times New Roman для всех текстовых элементов
    for item in ([ax.title, ax.xaxis.label, ax.yaxis.label] + 
                 ax.get_xticklabels() + ax.get_yticklabels()):
        item.set_fontname('Times New Roman')
        item.set_fontsize(10)
        item.set_color('black')
    
    plt.tight_layout()
    
    # Сохраняем график в буфер
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    buf.seek(0)
    
    # Добавляем график в документ
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(buf, width=Inches(6))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph("Рисунок: Зависимость динамического модуля упругости от удельного давления (средние значения)", style='Caption')
    apply_times_new_roman_12_black_to_paragraph(caption)

    buf.close()
    
    # Закрываем figure чтобы освободить память
    plt.close(fig)
    
    # 5. ДОБАВЛЯЕМ СРАВНИТЕЛЬНЫЕ ГРАФИКИ
    create_comparison_plots(samples_data, doc)
    
    # 6. Графики всех образцов с правильными подписями
    doc.add_page_break()
    heading = doc.add_heading('Графики всех образцов', level=1)
    apply_times_new_roman_12_black(heading)
    
    # Счетчик для нумерации рисунков
    figure_counter = 1
    
    for sample_id, data in samples_data.items():
        if 'images' not in data:
            continue
            
        heading = doc.add_heading(f'Образец {sample_id}', level=2)
        apply_times_new_roman_12_black(heading)
        
        # Получаем список масс для этого образца
        sample_masses = sorted([float(mass) for mass in data['results'].keys()]) if 'results' in data else []
        
        for image_name, fig in data['images'].items():
            # Извлекаем информацию о массе из имени файла
            mass_from_filename = None
            if '_' in image_name and 'кг' in image_name:
                # Пытаемся извлечь массу из имени файла (формат: "2_10.0кг.png")
                try:
                    # Разделяем по '_' и берем вторую часть
                    mass_part = image_name.split('_')[1]
                    # Удаляем расширение файла и "кг"
                    mass_str = mass_part.replace('кг', '').replace('.png', '').replace('.jpg', '')
                    mass_from_filename = float(mass_str)
                except (ValueError, IndexError):
                    mass_from_filename = None
            
            # Если не удалось извлечь из имени файла, берем из данных образца
            if mass_from_filename is None and sample_masses:
                # Берем первую массу из списка (можно адаптировать логику под вашу структуру)
                mass_from_filename = sample_masses[0] if figure_counter <= len(sample_masses) else sample_masses[-1]
            
            # Создаем подпись к рисунку
            if mass_from_filename is not None:
                caption_text = f"Рисунок {figure_counter}. Амплитудно-частотные характеристики образца {sample_id} при величине статического пригруза {mass_from_filename} кг. Модуль передаточной функции и график эффективности виброизоляции"
            else:
                caption_text = f"Рисунок {figure_counter}. Амплитудно-частотные характеристики образца {sample_id}. Модуль передаточной функции и график эффективности виброизоляции"
            
            # Подпись к рисунку
            caption = doc.add_paragraph(caption_text, style='Caption')
            apply_times_new_roman_12_black_to_paragraph(caption)
            
            # Вставка изображения
            buf = io.BytesIO()
            fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(buf, width=Inches(6))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            buf.close()
            
            # Увеличиваем счетчик рисунков
            figure_counter += 1
            
            doc.add_paragraph()  # Добавляем отступ между графиками
        
        doc.add_paragraph()  # Добавляем отступ между разными образцами
    
    doc.save(output_file)
    print(f"Полный отчет сохранен в: {output_file}")


def apply_times_new_roman_12_black(paragraph):
    """
    Применяет шрифт Times New Roman 12pt черный к заголовку
    """
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = None  # Черный цвет

def apply_times_new_roman_12_black_to_paragraph(paragraph):
    """
    Применяет шрифт Times New Roman 12pt черный к параграфу
    """
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = None  # Черный цвет

def apply_times_new_roman_black_to_cell(cell):
    """
    Применяет шрифт Times New Roman 12pt черный к ячейке таблицы
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = None  # Черный цвет

def convert_rar_to_zip(rar_path, zip_path):
    """Конвертирует RAR архив в ZIP архив"""
    try:
        with rarfile.RarFile(rar_path, 'r') as rar_ref:
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for file_info in rar_ref.infolist():
                    with rar_ref.open(file_info) as rar_file:
                        zip_ref.writestr(file_info.filename, rar_file.read())
        return True
    except Exception as e:
        print(f"Ошибка конвертации RAR в ZIP: {e}")
        return False

def extract_archive(archive_path, extract_to=None):
    """
    Извлекает архив (RAR или ZIP) во временную директорию
    """
    if extract_to is None:
        extract_to = tempfile.mkdtemp()
    
    os.makedirs(extract_to, exist_ok=True)
    
    archive_ext = os.path.splitext(archive_path)[1].lower()
    
    try:
        if archive_ext == '.zip':
            with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                zip_ref.extractall(extract_to)
                
        elif archive_ext == '.rar':
            # Создаем временный ZIP файл
            temp_zip = tempfile.mktemp(suffix='.zip')
            try:
                # Конвертируем RAR в ZIP
                if convert_rar_to_zip(archive_path, temp_zip):
                    # Извлекаем ZIP
                    with zipfile.ZipFile(temp_zip, 'r') as zip_ref:
                        zip_ref.extractall(extract_to)
                else:
                    raise Exception("Не удалось конвертировать RAR в ZIP")
            finally:
                # Удаляем временный ZIP файл
                if os.path.exists(temp_zip):
                    os.remove(temp_zip)
        else:
            raise ValueError(f"Unsupported archive format: {archive_ext}")
            
        return extract_to
        
    except Exception as e:
        # Очищаем временную директорию при ошибке
        if extract_to.startswith(tempfile.gettempdir()):
            shutil.rmtree(extract_to, ignore_errors=True)
        raise Exception(f"Failed to extract archive: {str(e)}")


def process_excel_file(excel_file):
    """
    Обрабатывает Excel файл с геометрическими параметрами и массами.
    Возвращает словарь с данными для каждого образца.
    
    Ожидаемая структура:
    Образец	Высота	Ширина	Длина	Масса	2	5	10
    1	    10	    100	    100	    120	    9	    8	    7
    2	    10	    100	    100	    125	    9	    8	    7
    
    Столбцы с массами (2, 5, 10) могут меняться.
    """
    
    try:
        # Читаем Excel файл
        df = pd.read_excel(excel_file)
        
        # Определяем основные колонки с фиксированными названиями
        fixed_columns = ['Образец', 'Высота', 'Ширина', 'Длина', 'Масса']
 
        # Находим колонки с массами (числовые значения)
        mass_columns = []
        for col in df.columns:
            # Пропускаем фиксированные колонки
            if col in fixed_columns:
                continue
            # Проверяем, является ли название колонки числом (массой)
            if str(col).replace('.', '').isdigit():
                mass_columns.append(col)
            # Также проверяем шаблон типа "2кг", "5_kg" и т.д.
            elif re.match(r'^\d+[\s_]*[kк]?[gг]?$', str(col), re.IGNORECASE):
                # Извлекаем числовую часть
                mass_value = re.findall(r'\d+', str(col))[0]
                mass_columns.append(mass_value)
 
        # Сортируем массы по возрастанию
        mass_columns = sorted([float(mass) for mass in mass_columns])
        mass_columns = [str(mass) for mass in mass_columns]
        samples_data = {}
        
        for _, row in df.iterrows():
            sample_id = str(int(row['Образец']))
            # Основные геометрические параметры
            geometric_params = {
                'height': float(row['Высота']),
                'width': float(row['Ширина']),
                'length': float(row['Длина']),
                'base_mass': float(row['Масса']) if 'Масса' in row else 0.0
            }
            
            # Данные по массам (пригрузам)
            mass_data = {}
            path_files = []
            for mass_col in mass_columns:
                mass_col = str(mass_col).split('.')[0]
                mass_col = int(mass_col)
                path_file = str(sample_id)+'_'+str(mass_col)+'.csv'
                path_files.append(path_file)
                if mass_col in row:
                    mass_value = float(mass_col)
                    # Здесь можно хранить дополнительные данные, например:
                    mass_data[mass_value] = {
                        'value': float(row[mass_col]) if pd.notna(row[mass_col]) else None,
                        # 'some_other_parameter': ...
                    }
            
            samples_data[sample_id] = {
                'geometric_params': geometric_params,
                'masses': mass_data,
                'all_mass_values': list(mass_data.keys()),  # список всех масс для этого образца
                'name_files': path_files
            }
        
        return samples_data, mass_columns

    except Exception as e:
        print(f"Ошибка при обработке Excel файла: {e}")

def extract_archive_to_temp(archive_path, extract_to=None):
    """
    Извлекает архив во временную директорию
    """
    if extract_to is None:
        extract_to = tempfile.mkdtemp()
    
    archive_ext = archive_path.split('.')[-1].lower()
    
    try:
        if archive_ext == 'zip':
            with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                zip_ref.extractall(extract_to)
        
        elif archive_ext == 'rar':
            with rarfile.RarFile(archive_path, 'r') as rar_ref:
                rar_ref.extractall(extract_to)
        
        return extract_to
        
    except Exception as e:
        if os.path.exists(extract_to):
            shutil.rmtree(extract_to, ignore_errors=True)
        raise Exception(f"Ошибка извлечения архива: {str(e)}")




def get_archive_files_list(archive_path):
    """
    Получает список файлов в архиве
    
    Args:
        archive_path: путь к архиву
    
    Returns:
        list: список имен файлов в архиве
    """
    archive_ext = archive_path.split('.')[-1].lower()
    files_list = []
    
    try:
        if archive_ext == 'zip':
            with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                files_list = zip_ref.namelist()
        
        elif archive_ext == 'rar':
            with rarfile.RarFile(archive_path, 'r') as rar_ref:
                files_list = [file.filename for file in rar_ref.infolist()]
    
    except Exception as e:
        raise Exception(f"Ошибка чтения архива: {str(e)}")
    
    return files_list


def get_file(temp_path, data): 
    list_files = []
    for i in data:
        path = (temp_path + '/' + i)
        list_files.append(path)
    return list_files

# Пример использования:
def example_usage():
    """
    Пример использования функции
    """
    # Обработка Excel файла
    samples_data, mass_columns = process_excel_file('Красный/Красный.xlsx')
    
    temp_path = (extract_archive('Красный/Красный.zip'))
 
    list_file = (get_archive_files_list('Красный/Красный.zip'))

    for sample_id, data in samples_data.items():
        for i in data['name_files']:
            if i  in list_file:
                continue
            else:
                print(f"Отсутствует файл {i} в архиеве")


        files = data['name_files'] 
        list_files = get_file(temp_path,files)
 
        a = data['geometric_params']['length']
        b = data['geometric_params']['width']
        h = data['geometric_params']['height']

        loads = mass_columns


        heights = [item['value'] for item in data['masses'].values()]
        

        images, datas, results = vibraTableOne(sample_id, list_files, a, b, h, heights, loads)

        samples_data[sample_id]['images'] = images
        samples_data[sample_id]['datas'] = datas
        samples_data[sample_id]['results'] = results

    create_full_report(samples_data, 'Красный/Красный.docx')


if __name__ == "__main__":
    example_usage()