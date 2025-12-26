import matplotlib
matplotlib.use('Agg')  # Устанавливаем неинтерактивный бэкенд
import math
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import datetime
import tempfile
import zipfile
from django.http import HttpResponse
import shutil
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
import base64
from scipy.signal import find_peaks
import numpy as np
from scipy import interpolate
import os
from docx import Document


# Настройки графиков
plt.rcParams['font.family'] = 'Times New Roman'  # Установка шрифта Times New Roman
plt.rcParams['font.style'] = 'normal'  # Обычный шрифт (не italic)
plt.rcParams['font.size'] = 12  # Размер шрифта 12
fontsize = 12
linewidth = 2
fontweight = 'bold'  # Обычный шрифт (не bold)

def save_plot(fig, filename):
    """Сохраняет график в файл и закрывает фигуру"""
    # Создаем временную директорию для графиков, если ее нет
    temp_dir = os.path.join(tempfile.gettempdir(), 'elastic_modulus_plots')
    os.makedirs(temp_dir, exist_ok=True)
    
    full_path = os.path.join(temp_dir, filename)
    fig.savefig(full_path, dpi=300, bbox_inches='tight')
    plt.close(fig)
    return full_path

def load_data(file_path):
    """Загружает данные из файла"""
    try:
        df = pd.read_csv(file_path, sep="\t", header=None)
        df.replace(",", ".", regex=True, inplace=True)
        df = df.astype(float)
        return df
    except Exception as e:
        print(f"Ошибка: Не удалось загрузить файл:\n{str(e)}")
        return None

def data_modul_young(df, width, length, height):
    """Вычисляет данные для построения графиков модуля упругости"""
    if df is None:
        print("Данные не загружены!")
        return None
    
    width = float(width)
    length = float(length)
    initial_height = float(height)
    area = width * length

    k = np.argmax(df[0].values > 0)
    M = df.values[k:, :4] - df.values[k, :4]
    sr = len(M) / M[-1, 3] if M[-1, 3] != 0 else 10

    F = M[:, 0]
    S = M[:, 2]
    T = np.arange(len(F)) / sr

    peaks, _ = find_peaks(S, height=0.5 * np.max(S))
    if len(peaks) < 3:
        print("Недостаточно пиков для анализа (минимум 3)")
        return None

    Start = peaks[2] - peaks[0] + 1
    Finish = peaks[2] + 1

    print('Start', Start)
    print('Finish', Finish)
    F1 = F[Start-1:Finish]
    S1 = S[Start-1:Finish]
    w = int(np.ceil(sr * 2))
    n = len(F1) // w
    print(w, 'Ширина окна')
    print(n, 'Количество циклов')
    # Точный аналог MATLAB кода
    Start = peaks[2] - peaks[0]  # locs(3)-locs(1)+1 (индексация с 0 в Python)
    Finish = peaks[2]  # locs(3)
    F1 = F[Start:Finish]  # F(Start:Finish)
    S1 = S[Start:Finish]  # S(Start:Finish)
    w = 2 * math.ceil(sr)  # w=2*ceil(sr)
    n = math.floor(len(F1) / w)  # floor(length(F1)/w)
    Pr = np.zeros(n)
    E1 = np.zeros(n) 
    Eps1 = np.zeros(n)
    print(f"Размер F1: {len(F1)}, w: {w}, n: {n}")
    for i in range(n):
        # Точная MATLAB индексация (начинается с 1)
        idx1 = i * w  # (i-1)*w+1 в MATLAB → i*w в Python (т.к. i начинается с 0)
        idx2 = (i + 1) * w - 1  # i*w в MATLAB → (i+1)*w-1 в Python
        
        # Проверяем границы (в MATLAB нет проверки, но добавим для безопасности)
        if idx2 >= len(F1):
            idx2 = len(F1) - 1
        
        # Точные MATLAB формулы:
        Pr[i] = (F1[idx1] + F1[idx2]) / (2 * area)  # (F1((i-1)*w+1)+F1(i*w))/2/A
        
        delta_F = F1[idx2] - F1[idx1]  # F1(i*w)-F1((i-1)*w+1)
        delta_S = S1[idx2] - S1[idx1]  # S1(i*w)-S1((i-1)*w+1)
        
        # E1(i)=(F1(i*w)-F1((i-1)*w+1))./A./((S1(i*w)-S1((i-1)*w+1))./h0)
        if delta_S != 0:
            E1[i] = (delta_F / area) / (delta_S / initial_height)
        else:
            E1[i] = 0
        
        # Eps1(i)=(S1((i-1)*w+1)+S1(i*w))/2/h0
        Eps1[i] = (S1[idx1] + S1[idx2]) / (2 * initial_height)
    # Точный аналог MATLAB коррекции нуля
    if len(Pr) > 0:
        del_val = Pr[0]  # del=Pr(1)
        for i in range(len(Pr)):
            Pr[i] = Pr[i] - del_val  # Pr(i)=Pr(i)-del
    # Убираем лишние преобразования которые были в старом коде
    Eps1 = Eps1 * 100  # преобразование в %
    E1 = E1  # оставляем как есть (в МПа)
    Pr = Pr  # оставляем как есть (в МПа)
    print(f"Результат: Pr[{len(Pr)}], E1[{len(E1)}], Eps1[{len(Eps1)}]")
    
    return E1, Eps1, Pr

def create_plot_modul_young(E1, Eps1, Pr, name_sample, form_factor):
    """Создает график модуля Юнга"""
    fig, (ax4, ax5) = plt.subplots(nrows=2, ncols=1, figsize=(10, 6), sharex=True)

    ax4.plot(Pr, E1, 'k-', linewidth=linewidth)
    ax4.set_title(f'{name_sample} | q = {form_factor:.4f}', 
                 fontsize=fontsize, fontweight=fontweight)
    ax4.set_xlabel('Удельное давление, МПа', fontsize=fontsize, fontweight=fontweight)
    ax4.set_ylabel('Модуль упругости, МПа', fontsize=fontsize, fontweight=fontweight)
    ax4.grid(True, linestyle='--', alpha=0.6)
    ax4.set_xlim(left=0)
    ax4.set_ylim(bottom=0)

    ax5.plot(Pr, Eps1, 'k-', linewidth=linewidth)
    ax5.set_xlabel('Удельное давление, МПа', fontsize=fontsize, fontweight=fontweight)
    ax5.set_ylabel('Относительная деформация, %', fontsize=fontsize, fontweight=fontweight)
    ax5.grid(True, linestyle='--', alpha=0.6)
    ax5.set_xlim(left=0)
    ax5.set_ylim(bottom=0)

    plt.tight_layout()
    plt.close()
    return save_plot(fig, f"{name_sample}_modul_young.png")


def full_plot(Time, Disp, Forse, name_sample, form_factor):
    """Создает полный график зависимости перемещения и нагрузки от времени"""
    fig, ax1 = plt.subplots(figsize=(10, 6))
    
    # Инициализация возвращаемых значений по умолчанию
    unload_time = None
    unload_force = None
    last_peak_force = None
    time_load = None
    delta_force = None
    
    # Построение графика перемещения
    ax1.plot(Time, Disp, 'b', label='Смещение (мм)', linewidth=linewidth)
    ax1.set_ylabel('Перемещение, мм', fontsize=fontsize, fontweight=fontweight, color='blue')
    ax1.set_xlabel('Время, с', fontsize=fontsize, fontweight=fontweight)
    ax1.grid(True)
    ax1.set_ylim([0, math.ceil(max(Disp) + 0.5)])

    # Построение графика нагрузки
    ax1_force = ax1.twinx()
    line, = ax1_force.plot(Time, Forse, 'r', label='Нагрузка (Н)', linewidth=linewidth)
    ax1_force.set_ylabel('Нагрузка, Н', fontsize=fontsize, fontweight=fontweight, color='red')
    ax1.set_title(f'{name_sample} | q = {form_factor:.4f}', 
                 fontsize=fontsize, fontweight=fontweight)

    # Находим пики нагрузки
    peaks, _ = find_peaks(Forse, prominence=np.std(Forse)/2)
    
    if len(peaks) > 0:
        # Берем последний пик
        last_peak_idx = peaks[-1]
        last_peak_time = Time[last_peak_idx]
        last_peak_force = Forse[last_peak_idx]
        
        # Анализируем перемещение после пика
        disp_after_peak = Disp[last_peak_idx:]
        
        # Находим точку, где перемещение начинает резко уменьшаться
        disp_derivative = np.gradient(disp_after_peak)
        
        # Ищем точку разгрузки
        threshold = -0.001
        unload_idx = last_peak_idx
        for i in range(1, len(disp_derivative)):
            if disp_derivative[i] < threshold:
                unload_idx = last_peak_idx + i
                break
        
        unload_time = Time[unload_idx]

        # Конечная нагрузка 
        unload_force = Forse[unload_idx]

        # Время релаксации 
        time_load = unload_time - last_peak_time
        
        delta_force = last_peak_force - unload_force
        # Визуализация
        if time_load >= 500:
            ax1_force.plot(last_peak_time, last_peak_force, 'go', markersize=8, label='Пик нагрузки')
            ax1_force.plot(unload_time, unload_force, 'mo', markersize=8, label='Начало разгрузки')
        

    print(name_sample)
    print(last_peak_force - unload_force)
    print(time_load)
    print('#' * 25)
    print(delta_force, 'delta_force')


    plt.tight_layout()
    plt.close()
    plot_path = save_plot(fig, f"{name_sample}_full_plot.png")
    
    # Возвращаем путь к графику и нужные значения
    return {
        'plot_path': plot_path,
        'unload_time': unload_time,
        'unload_force': unload_force, # Нагрузка в конце релаксации
        'last_peak_force': last_peak_force, # Нагрузка на последнем пике в начале релаксации 
        'time_load': time_load,
        'delta_force': delta_force
    }


def plot_cycles_only(Forse, Disp, locs, name_sample, form_factor):
    """Создает графики циклов нагружения"""
    if len(locs) < 1:
        print(f"Недостаточно данных для построения графика {name_sample}")
        return None

    def colors_labels(n):
        base_colors = ['k', 'g', 'r', 'b', 'm', 'c', 'y']
        n = min(n, len(base_colors))
        labels = [f'Цикл {i+1}' for i in range(n)]
        return base_colors[:n], labels

    n_locs_rang = min(6, len(locs))
    colors, labels = colors_labels(n_locs_rang)
    base_len = locs[0] if len(locs) > 0 else len(Disp)

    fig, ax = plt.subplots(figsize=(10, 6))
    for i in range(n_locs_rang):
        start = max(locs[i] - base_len + 1, 0)
        end = min(locs[i] + base_len, len(Disp))
        ax.plot(Disp[start:end], Forse[start:end], colors[i], label=labels[i], linewidth=linewidth)

    ax.set_xlabel('Перемещение, мм', fontsize=fontsize, fontweight=fontweight)
    ax.set_ylabel('Нагрузка, Н', fontsize=fontsize, fontweight=fontweight)
    ax.grid(True)
    ax.legend(loc='lower right')
    ax.set_title(f'{name_sample} | q = {form_factor:.4f}',
                fontsize=fontsize, fontweight=fontweight)

    plt.tight_layout()
    plt.close()
    return save_plot(fig, f"{name_sample}_cycles.png")

def find_loading_cycles(df):
    """Находит циклы нагружения в данных"""
    peaks, _ = find_peaks(df[0].values, prominence=np.std(df[0].values)/2)
    return peaks.tolist()

def process_sample_file(filepath, sample_name, width, length, height, mass):
    """Обрабатывает данные одного образца и создает графики"""

    df = load_data(filepath)
    if df is None:
        return None

    # Нахождения площади образца
    A = width * length

    # Нахождение коэффициента формы
    q = A / (2 * height * ( width + length))

    force = df[0].values
    displacement = df[2].values
    time = df[3].values
    locs = find_loading_cycles(df)

    plot_data = full_plot(time, displacement, force, sample_name, q)

    full_plot_path = plot_data['plot_path']

    unload_time = plot_data['unload_time']
    unload_force = plot_data['unload_force']
    last_peak_force = plot_data['last_peak_force']
    time_load = plot_data['time_load']
    delta_force = plot_data['delta_force']

    E1, Eps1, Pr = data_modul_young(df, width, length, height)
    
    if E1 is None:
        return None
        
    modul_plot_path = create_plot_modul_young(E1, Eps1, Pr, sample_name, q)
    cycles_plot_path = plot_cycles_only(force, displacement, locs, sample_name, q)
    
    # Получаем значения нагрузки в заданных диапазонах деформации
    load_values = get_load_at_deformations(Pr, Eps1)

    return {
        'name': sample_name,
        'width': width,
        'length': length,
        'height': height,
        'mass': mass,
        'full_plot': full_plot_path,
        'modul_plot': modul_plot_path,
        'cycles_plot': cycles_plot_path if cycles_plot_path else None,
        'load_values': load_values,
        'unload_time': unload_time,
        'unload_force': unload_force,
        'last_peak_force': last_peak_force,
        'time_load': time_load,
        'delta_force': delta_force,
        'E1':E1,
        'Eps1':Eps1,
        'Pr': Pr
    }

def insert_load_table(doc, samples_data, decimal_places=3):
    """Создает таблицу с нагрузками при разных деформациях"""
    if not samples_data or 'load_values' not in samples_data[0]:
        return
    
    # Определяем заголовки на основе первого образца
    first_sample = samples_data[0]
    headers = ["Образец"] + list(first_sample['load_values'].keys())
    
    table = doc.add_table(rows=len(samples_data)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header.replace("max_deformation", "макс. деформация"))
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Данные образцов
    for row_idx, sample in enumerate(samples_data, start=1):
        row_cells = table.rows[row_idx].cells
        
        # Формируем строку данных
        data = [sample['name']]
        for key in headers[1:]:
            value = sample['load_values'].get(key, 0.0)
            if "max_deformation" in key:  # Для случая максимальной деформации
                formatted_value = f"{value:.{decimal_places}f}".replace('.', ',')
            else:
                formatted_value = f"{value:.{decimal_places}f}".replace('.', ',')
            data.append(formatted_value)
        
        for col_idx, value in enumerate(data):
            p = row_cells[col_idx].paragraphs[0]
            run = p.add_run(value)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Добавляем подпись к таблице
    p = doc.add_paragraph()
    run = p.add_run("Таблица 1 - Величина удельной нагрузки при относительной деформации, Н/мм²")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Добавляем пустой абзац для отступа


def get_load_at_deformations(Pr, Eps1, deformation_ranges=[(10, 11), (19, 21), (38, 40)]):
    """Возвращает значения нагрузки при указанных диапазонах деформации"""
    results = {}
    
    for low, high in deformation_ranges:
        # Находим все точки, где деформация попадает в диапазон
        mask = (Eps1 >= low) & (Eps1 <= high)
        indices = np.where(mask)[0]
        
        if len(indices) > 0:
            # Среднее значение нагрузки в диапазоне
            avg_load = np.mean(Pr[indices])
            results[f"{low}-{high}%"] = avg_load
        else:
            # Если нет данных, ищем ближайшее значение
            closest_idx = np.argmin(np.abs(Eps1 - (low + high)/2))
            results[f"{low}-{high}%"] = Pr[closest_idx]
    
    return results


def add_custom_heading(doc, text, level=1):
    """Добавляет заголовок с ручным форматированием"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    
    # Настройки форматирования в зависимости от уровня заголовка
    if level == 1:
        run.bold = True
        run.font.size = Pt(16)
    elif level == 2:
        run.bold = True
        run.font.size = Pt(14)
    else:
        run.bold = True
        run.font.size = Pt(12)
    return p


def insert_samples_table(doc, samples_data, decimal_places={'length': 2, 'width': 2, 'height': 2, 'mass': 2}):
    """Создает таблицу с параметрами образцов с настраиваемым количеством знаков после запятой
    
    Args:
        doc: Объект документа Word
        samples_data: Данные образцов
        decimal_places: Словарь с количеством знаков после запятой для каждого параметра
            По умолчанию: длина/ширина/высота/ масса - 2 знака
    """
    table = doc.add_table(rows=len(samples_data)+1, cols=5)
    table.style = 'Table Grid'
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    headers = ["Образец", "Длина, мм", "Ширина, мм", "Толщина, мм", "Масса, г"]
    
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Данные образцов с форматированием
    for row_idx, sample in enumerate(samples_data, start=1):
        row_cells = table.rows[row_idx].cells
        
        # Форматируем каждое значение с нужным количеством знаков
        length_str = f"{float(sample['length']):.{decimal_places['length']}f}".replace('.', ',')
        width_str = f"{float(sample['width']):.{decimal_places['width']}f}".replace('.', ',')
        height_str = f"{float(sample['height']):.{decimal_places['height']}f}".replace('.', ',')
        mass_str = f"{float(sample['mass']):.{decimal_places['mass']}f}".replace('.', ',')
        
        data = [
            sample['name'],
            length_str,
            width_str,
            height_str,
            mass_str
        ]
        
        for col_idx, value in enumerate(data):
            p = row_cells[col_idx].paragraphs[0]
            run = p.add_run(value)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def insert_samples_graphs(doc, samples_data):
    """Вставляет графики с подписями"""
    figure_counter = 3
    
    for sample in samples_data:
        graphs = [
            ('full_plot', f"График зависимости перемещения и нагрузки от времени образца {sample['name']}"),
            ('modul_plot', f"График модуля упругости образца {sample['name']}"),
            ('cycles_plot', f"Графики циклов нагружения образца {sample['name']}")
        ]
        
        for plot_type, description in graphs:
            if sample.get(plot_type) and os.path.exists(sample[plot_type]):
                # Добавляем разрыв страницы перед каждым графиком (кроме первого)
                if figure_counter > 1:
                    doc.add_page_break()
                
                # Вставляем график
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(sample[plot_type], width=Inches(9))  # Чуть меньший размер для лучшего отображения
                
                # Добавляем подпись
                p = doc.add_paragraph()
                run = p.add_run(f"Рисунок {figure_counter}. {description}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                figure_counter += 1


def fill_template(template_path, samples_data, output_filename):
    """Заполняет шаблон документа данными с правильным расположением элементов"""
    doc = Document(template_path)
    
    # Сначала обрабатываем все параграфы шаблона
    for paragraph in list(doc.paragraphs):  # Используем list() для создания копии
        text = paragraph.text

        print(text)
        
        if '{ДАТА}' in text:
            # Заменяем дату
            paragraph.text = text.replace('{ДАТА}', datetime.datetime.now().strftime("%d.%m.%Y"))
            
        elif '{LOAD_TABLE_PLACEHOLDER}' in text:
            # Вставляем таблицу нагрузок вместо плейсхолдера
            paragraph.text = paragraph.text.replace('{LOAD_TABLE_PLACEHOLDER}', '')
            insert_load_table(doc, samples_data)
            
        elif '{SAMPLES_TABLE_PLACEHOLDER}' in text:
            # Вставляем таблицу параметров образцов вместо плейсхолдера
            paragraph.text = paragraph.text.replace('{SAMPLES_TABLE_PLACEHOLDER}', '')
            insert_samples_table(doc, samples_data)
            
        elif '{GRAPHS_PLACEHOLDER}' in text:
            # Вставляем графики вместо плейсхолдера
            paragraph.text = paragraph.text.replace('{GRAPHS_PLACEHOLDER}', '')
            
            # Добавляем заголовок перед графиками
            p = doc.add_paragraph()
            run = p.add_run("Графики результатов испытаний")
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Вставляем сами графики
            insert_samples_graphs(doc, samples_data)
    
    doc.save(output_filename)



def fill_template_(template_path, samples_data, output_filename):
    """Заполняет шаблон документа данными с правильным расположением элементов"""
    try:
        doc = Document(template_path)

        # Сначала обрабатываем все параграфы шаблона
        for paragraph in list(doc.paragraphs):  # Используем list() для создания копии
            text = paragraph.text

            if '{ДАТА}' in text:
                # Заменяем дату
                paragraph.text = text.replace('{ДАТА}', datetime.datetime.now().strftime("%d.%m.%Y"))
                
            elif '{LOAD_TABLE_PLACEHOLDER}' in text:
                # Вставляем таблицу нагрузок вместо плейсхолдера
                paragraph.text = paragraph.text.replace('{LOAD_TABLE_PLACEHOLDER}', '')
                insert_load_table(doc, samples_data)
                
            elif '{SAMPLES_TABLE_PLACEHOLDER}' in text:
                # Вставляем таблицу параметров образцов вместо плейсхолдера
                paragraph.text = paragraph.text.replace('{SAMPLES_TABLE_PLACEHOLDER}', '')
                insert_samples_table(doc, samples_data)
                
            elif '{GRAPHS_PLACEHOLDER}' in text:
                # Вставляем графики вместо плейсхолдера
                paragraph.text = paragraph.text.replace('{GRAPHS_PLACEHOLDER}', '')
                
                # Добавляем заголовок перед графиками
                p = doc.add_paragraph()
                run = p.add_run("Графики результатов испытаний")
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Вставляем сами графики
                insert_samples_graphs(doc, samples_data)
        
        # Сохраняем документ
        doc.save(output_filename)
        return True
    
    except Exception as e:
        print(f"Ошибка при заполнении шаблона: {str(e)}")
        return False


def cleanup_temp_files(samples_data):
    """Удаляет временные файлы с графиками"""
    for sample in samples_data:
        for plot_type in ['full_plot', 'modul_plot', 'cycles_plot']:
            if sample.get(plot_type) and os.path.exists(sample[plot_type]):
                try:
                    os.remove(sample[plot_type])
                except:
                    pass


def genaretion_plot(data_list, data_excel, template_path=None, output_filename='Итоговый_протокол.docx'):
        """
        Функция для создания объедененного протокола статика.
        Функция проходиться по файлу добавляет информацию по тегам {NAME_TAG}.
            
            :param data_list: список с испытаниями 
            :param data_excel: Excel файл с размерами образца
            :param template_path: путь до шаблона Word-файла
            :param output_filename: конечное название файла
        """

        # Определяем путь к шаблону по умолчанию
        if template_path is None:
            template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
            if not os.path.exists(template_path):
                # Создаем пустой документ, если шаблона нет
                doc = Document()
                doc.add_paragraph('{ДАТА}')
                doc.add_paragraph('{LOAD_TABLE_PLACEHOLDER}')
                doc.add_paragraph('{SAMPLES_TABLE_PLACEHOLDER}')
                doc.add_paragraph('{GRAPHS_PLACEHOLDER}')
                doc.save(template_path)
        
        # Проверяем наличие всех плейсхолдеров в шаблоне
        required_placeholders = {'{ДАТА}', '{LOAD_TABLE_PLACEHOLDER}', 
                               '{SAMPLES_TABLE_PLACEHOLDER}', '{GRAPHS_PLACEHOLDER}'}
        doc = Document(template_path)
        existing_placeholders = set()
        for paragraph in doc.paragraphs:
            for placeholder in required_placeholders:
                if placeholder in paragraph.text:
                    existing_placeholders.add(placeholder)
        
        missing_placeholders = required_placeholders - existing_placeholders
        if missing_placeholders:
            print(f"Внимание: В шаблоне отсутствуют следующие плейсхолдеры: {', '.join(missing_placeholders)}")
        
        # Остальная часть функции остается без изменений
        data_excel['Образец'] = data_excel['Образец'].astype(str).str.strip()
        data_excel.columns = data_excel.columns.str.strip()
        samples_data = []
        
        for filepath in data_list:
            filename = os.path.basename(filepath)
            sample_name = os.path.splitext(filename)[0]
            
            row = data_excel[data_excel['Образец'] == sample_name]
            if row.empty:
                print(f"Образец {sample_name} не найден в таблице!")
                continue

            try:
                # Convert numpy.float64 to float directly (no need for replace)
                width = float(row['Ширина'].values[0].replace(',', '.'))
                length = float(row['Длина'].values[0].replace(',', '.'))
                height = float(row['Высота'].values[0].replace(',', '.'))
                mass = float(row['Масса'].values[0].replace(',', '.'))
            except (IndexError, ValueError) as e:
                print(f"Ошибка получения параметров для образца {sample_name}: {str(e)}")
                continue

            sample_data = process_sample_file(filepath, sample_name, width, length, height, mass)
            if sample_data:
                samples_data.append(sample_data)
        
        if samples_data:
            fill_template(template_path, samples_data, output_filename)
            cleanup_temp_files(samples_data)
            print(f"Протокол успешно сохранен в файл: {output_filename}")
            return True
        else:
            print("Нет данных для создания протокола!")
            return False

def save_sample_data_to_excel(sample_data, output_dir):
    """
    Сохраняет данные образца (E1, Eps1, Pr) в Excel файл с интерполяцией и релаксацией
    
    Args:
        sample_data: словарь с данными образца
        output_dir: директория для сохранения файла
        
    Returns:
        Путь к сохраненному Excel файлу или None в случае ошибки
    """
    try:
        sample_name = sample_data['name']
        
        # Проверяем наличие основных данных
        required_data = ['E1', 'Eps1', 'Pr']
        for data_key in required_data:
            if data_key not in sample_data:
                print(f"Нет данных {data_key} для образца {sample_name}")
                return None
        
        # Извлекаем данные
        E1 = sample_data['E1']  # Модуль упругости (Estat), МПа
        Eps1 = sample_data['Eps1']  # Относительная деформация, %
        Pr = sample_data['Pr']  # Удельное давление, МПа
        
        # Проверяем, что данные не пустые
        if len(Eps1) == 0 or len(Pr) == 0 or len(E1) == 0:
            print(f"Пустые данные для образца {sample_name}")
            return None
        
        # Создаем DataFrame с исходными данными
        df_original = pd.DataFrame({
            'Удельное давление, МПа': Pr,
            'Относительная деформация, %': Eps1,
            'Estat, МПа': E1
        })
        
        # Добавляем интерполяцию для заданных процентов деформации
        df_interpolation = pd.DataFrame()
        if len(Eps1) > 1:  # Для интерполяции нужно минимум 2 точки
            percentages = np.arange(5, 51, 5)  # 5%, 10%, 15%, ..., 50%
            
            try:
                # Проверяем, что Eps1 отсортирован
                if not all(Eps1[i] <= Eps1[i+1] for i in range(len(Eps1)-1)):
                    # Сортируем данные по Eps1
                    sorted_indices = np.argsort(Eps1)
                    Eps1_sorted = np.array(Eps1)[sorted_indices]
                    Pr_sorted = np.array(Pr)[sorted_indices]
                    E1_sorted = np.array(E1)[sorted_indices]
                else:
                    Eps1_sorted = Eps1
                    Pr_sorted = Pr
                    E1_sorted = E1
                
                # Определяем реальные границы данных
                min_eps = float(min(Eps1_sorted))
                max_eps = float(max(Eps1_sorted))
                
                print(f"Образец {sample_name}: мин. деф. = {min_eps}%, макс. деф. = {max_eps}%")
                
                # Создаем интерполяционные функции ТОЛЬКО для интерполяции
                f_load = interpolate.interp1d(
                    Eps1_sorted, Pr_sorted, 
                    kind='linear', 
                    bounds_error=False,
                    fill_value=np.nan
                )
                
                # Получаем интерполированные значения
                interpolated_data = []
                
                for pct in percentages:
                    # Интерполируем ТОЛЬКО если процент в пределах измеренных данных
                    if min_eps <= pct <= max_eps:
                        load_interp = float(f_load(pct))
                        
                        # Проверяем, что получили числа (не NaN)
                        if not np.isnan(load_interp):
                            interpolated_data.append({
                                'Процент сжатия образца, %': pct,
                                'Интерполированное удельное давление, МПа': load_interp,
                            })
                        else:
                            print(f"  Предупреждение: не удалось интерполировать для {pct}%")
                    else:
                        print(f"  Пропуск {pct}%: вне диапазона измерений ({min_eps:.1f}%-{max_eps:.1f}%)")
                
                if interpolated_data:
                    df_interpolation = pd.DataFrame(interpolated_data)
                    
            except Exception as e:
                print(f"Ошибка интерполяции для образца {sample_name}: {str(e)}")
        
        # Создаем лист с данными релаксации
        df_relaxation = pd.DataFrame()
        # Проверяем наличие данных релаксации по ключам из возвращаемой структуры
        if ('unload_time' in sample_data and sample_data['unload_time'] is not None and 
            'unload_force' in sample_data and sample_data['unload_force'] is not None):

            # Проверка на то чтобы в таблицу попадали образцы с испытаниями на релаксацию
            if sample_data['time_load'] >= 500:
            
                relaxation_rows = []
                
                # Время релаксации (разгрузки)
                if sample_data['unload_time'] is not None:
                    relaxation_rows.append({
                        'Параметр': 'Время релаксации (t, с)',
                        'Значение': f"{sample_data['time_load']:.2f} с"
                    })
                
                # Начальная нагрузка (последний пик)
                if 'last_peak_force' in sample_data and sample_data['last_peak_force'] is not None:
                    relaxation_rows.append({
                        'Параметр': 'Начальная нагрузка',
                        'Значение': f"{sample_data['last_peak_force']:.2f} Н"
                    })
                
                # Конечная нагрузка
                if sample_data['unload_force'] is not None:
                    relaxation_rows.append({
                        'Параметр': 'Конечная нагрузка',
                        'Значение': f"{sample_data['unload_force']:.2f} Н"
                    })
                # Конечная нагрузка
                if sample_data['delta_force'] is not None:
                    relaxation_rows.append({
                        'Параметр': 'Изминение нагрузки (R, Н)',
                        'Значение': f"{sample_data['delta_force']:.2f} Н"
                    })
                
                if relaxation_rows:
                    df_relaxation = pd.DataFrame(relaxation_rows)
        
        # Путь для сохранения (ВНЕ всех условий, всегда определяется)
        excel_path = os.path.join(output_dir, f"{sample_name}_data.xlsx")
        
        # Сохраняем в Excel
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            # Исходные данные
            df_original.to_excel(writer, sheet_name='Исходные данные', index=False)
            
            # Интерполяция (если есть данные)
            if not df_interpolation.empty:
                df_interpolation.to_excel(writer, sheet_name='Интерполяция', index=False)
            
            # Релаксация (если есть данные)
            if not df_relaxation.empty:
                df_relaxation.to_excel(writer, sheet_name='Релаксация', index=False)
            
            # Настраиваем ширину колонок
            for sheet_name in writer.sheets:
                worksheet = writer.sheets[sheet_name]
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
        
        print(f"Данные сохранены в Excel: {excel_path}")
        return excel_path
        
    except Exception as e:
        print(f"Ошибка при сохранении Excel для образца {sample_name}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def genaretion_plot_with_saved_plots(data_list, data_excel, template_path=None, output_dir=None, output_filename='Итоговый_протокол.docx'):
    """
    Функция для создания объединенного протокола статика с сохранением графиков в отдельные папки
    и Excel файлов с данными.
    
    :param data_list: список с путями к файлам испытаний 
    :param data_excel: DataFrame с размерами образца
    :param template_path: путь до шаблона Word-файла
    :param output_dir: директория для сохранения выходных файлов
    :param output_filename: конечное название файла
    :return: (success, plot_dirs, excel_files)
    """
    
    if output_dir is None:
        output_dir = os.path.join(tempfile.gettempdir(), 'vibration_analysis')
    
    os.makedirs(output_dir, exist_ok=True)
    
    # Определяем путь к шаблону по умолчанию
    if template_path is None:
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        if not os.path.exists(template_path):
            # Создаем пустой документ, если шаблона нет
            doc = Document()
            doc.add_paragraph('{ДАТА}')
            doc.add_paragraph('{LOAD_TABLE_PLACEHOLDER}')
            doc.add_paragraph('{SAMPLES_TABLE_PLACEHOLDER}')
            doc.add_paragraph('{GRAPHS_PLACEHOLDER}')
            doc.save(template_path)
    
    # Проверяем наличие всех плейсхолдеров в шаблоне
    required_placeholders = {'{ДАТА}', '{LOAD_TABLE_PLACEHOLDER}', 
                           '{SAMPLES_TABLE_PLACEHOLDER}', '{GRAPHS_PLACEHOLDER}'}
    doc = Document(template_path)
    existing_placeholders = set()
    for paragraph in doc.paragraphs:
        for placeholder in required_placeholders:
            if placeholder in paragraph.text:
                existing_placeholders.add(placeholder)
    
    missing_placeholders = required_placeholders - existing_placeholders
    if missing_placeholders:
        print(f"Внимание: В шаблоне отсутствуют следующие плейсхолдеры: {', '.join(missing_placeholders)}")
    
    # Остальная часть функции остается без изменений
    data_excel['Образец'] = data_excel['Образец'].astype(str).str.strip()
    data_excel.columns = data_excel.columns.str.strip()
    samples_data = []
    plot_dirs = {}  # Словарь для хранения путей к папкам с графиками
    excel_files = []  # НОВОЕ: список для хранения путей к Excel файлам
    
    for filepath in data_list:
        filename = os.path.basename(filepath)
        sample_name = os.path.splitext(filename)[0]
        
        row = data_excel[data_excel['Образец'] == sample_name]
        if row.empty:
            print(f"Образец {sample_name} не найден в таблице!")
            continue

        try:
            # Convert numpy.float64 to float directly (no need for replace)
            width = float(row['Ширина'].values[0].replace(',', '.'))
            length = float(row['Длина'].values[0].replace(',', '.'))
            height = float(row['Высота'].values[0].replace(',', '.'))
            mass = float(row['Масса'].values[0].replace(',', '.'))
        except (IndexError, ValueError) as e:
            print(f"Ошибка получения параметров для образца {sample_name}: {str(e)}")
            continue

        # Создаем папку для графиков этого образца
        sample_plot_dir = os.path.join(output_dir, 'graphs', sample_name)
        os.makedirs(sample_plot_dir, exist_ok=True)
        plot_dirs[sample_name] = sample_plot_dir
        
        # Временная модификация функции save_plot для сохранения в нужную папку
        original_save_plot = save_plot
        
        def custom_save_plot(fig, filename):
            """Кастомная функция сохранения графиков в папку образца"""
            full_path = os.path.join(sample_plot_dir, filename)
            fig.savefig(full_path, dpi=300, bbox_inches='tight')
            plt.close(fig)
            return full_path
        
        # Временно заменяем функцию save_plot
        import protocol.utils.protocol as protocol_module
        protocol_module.save_plot = custom_save_plot
        
        try:
            # Вызываем оригинальную функцию process_sample_file
            sample_data = process_sample_file(filepath, sample_name, width, length, height, mass)
            
            if sample_data:
                # НОВОЕ: Сохраняем данные в Excel файл
                excel_path = save_sample_data_to_excel(sample_data, sample_plot_dir)
                if excel_path:
                    excel_files.append(excel_path)
                    sample_data['excel_file'] = excel_path
                
                # Обновляем пути к графикам в sample_data, чтобы они указывали на правильную папку
                plot_types = ['full_plot', 'modul_plot', 'cycles_plot']
                for plot_type in plot_types:
                    if sample_data.get(plot_type):
                        # Получаем имя файла из старого пути
                        old_path = sample_data[plot_type]
                        if old_path:
                            filename = os.path.basename(old_path)
                            # Создаем новый путь в папке образца
                            new_path = os.path.join(sample_plot_dir, filename)
                            # Обновляем путь в данных
                            sample_data[plot_type] = new_path
                
                samples_data.append(sample_data)
        finally:
            # Восстанавливаем оригинальную функцию
            protocol_module.save_plot = original_save_plot
    
    if samples_data:
        # Создаем протокол с помощью оригинальной функции fill_template
        fill_template(template_path, samples_data, output_filename)
        
        # НЕ удаляем временные файлы (графики нам нужны в архиве)
        # cleanup_temp_files(samples_data)
        
        print(f"Протокол успешно сохранен в файл: {output_filename}")
        print(f"Графики сохранены в директории: {os.path.join(output_dir, 'graphs')}")
        print(f"Excel файлов создано: {len(excel_files)}")  # НОВОЕ
        
        return True, plot_dirs, excel_files  # НОВОЕ: возвращаем excel_files
    else:
        print("Нет данных для создания протокола!")
        return False, {}, []  # НОВОЕ: возвращаем пустой список excel_files


def generate_individual_protocols(data_list, data_excel, template_path=None, output_dir='Индивидуальные_протоколы', zip_response=False):
    """
    Создает отдельные протоколы для каждого образца с возможностью архивации
    """
    # Подготовка данных
    data_excel['Образец'] = data_excel['Образец'].astype(str).str.strip()
    data_excel.columns = data_excel.columns.str.strip()
    
    with tempfile.TemporaryDirectory() as temp_dir:
        zip_buffer = BytesIO()
        created_files = []
        
        for filepath in data_list:
            filename = os.path.basename(filepath)
            sample_name = os.path.splitext(filename)[0]
            
            row = data_excel[data_excel['Образец'] == sample_name]
            if row.empty:
                print(f"Образец {sample_name} не найден в таблице!")
                continue

            def safe_convert(value):
                if isinstance(value, str):
                    return float(value.replace(',','.'))
                return float(value)
            
            # Получаем параметры образца
            width = safe_convert(row['Ширина'].values[0])
            length = safe_convert(row['Длина'].values[0])
            height = safe_convert(row['Высота'].values[0])
            mass = safe_convert(row['Масса'].values[0])
            protocol_number = row['Протокол'].values[0] if 'Протокол' in row.columns else sample_name
            
            # Обрабатываем данные образца
            sample_data = process_sample_file(filepath, sample_name, width, length, height, mass)
            if not sample_data:
                print(f"Не удалось обработать данные для образца {sample_name}")
                continue
            
            # Подготавливаем данные для шаблона
            template_data = {
                'length': length,
                'width': width,
                'height': height,
                'mass': mass,
                'name': sample_name,
                # Добавляем остальные поля из sample_data если они нужны в шаблоне
                **sample_data
            }
            
            
            print("Данные для шаблона:", template_data)
            
            # Создаем временный docx файл
            temp_output = os.path.join(temp_dir, f"Протокол_{protocol_number}.docx")
            
            # Заполняем шаблон
            success = fill_template_(template_path, template_data, temp_output)
            print("Результат заполнения шаблона:", success)
            
            if success:
                created_files.append(temp_output)
            else:
                print(f"Ошибка при заполнении шаблона для образца {sample_name}")
        
        # Формируем результат
        if zip_response and created_files:
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in created_files:
                    zipf.write(file_path, os.path.basename(file_path))
            
            zip_buffer.seek(0)
            response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="protocols_archive.zip"'
            return response
        
        elif not zip_response and created_files:
            os.makedirs(output_dir, exist_ok=True)
            for file_path in created_files:
                dest_path = os.path.join(output_dir, os.path.basename(file_path))
                shutil.move(file_path, dest_path)
            
            print(f"Создано {len(created_files)} протоколов в папке '{output_dir}'")
            return True
        

def save_plot_to_html(fig):
    """Сохраняет график matplotlib в HTML-совместимый формат"""
    buf = BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=100)
    plt.close(fig)
    buf.seek(0)
    image_base64 = base64.b64encode(buf.read()).decode('utf-8')
    return f"data:image/png;base64,{image_base64}"

def read_ecofizika(file, axes=['2','1']):
    """Reads data from Ecofizika (Octava)"""
    vibration = pd.read_csv(file, sep='\t', encoding='mbcs', header=None, names=axes,
                          dtype=np.float32,
                          skiprows=4, usecols=range(1,len(axes)+1)).reset_index(drop=True)
    inf = pd.read_csv(file, sep=' ', encoding='mbcs', header=None, names=None,
                     skiprows=2, nrows=1).reset_index(drop=True)
    fs = int(inf.iloc[0, -1])
    return vibration, fs

def find_res_width2(TR, freqs, peak_pos):
    """Нахождение ширины резонанса на половине высоты"""
    try:
        half_height = TR[peak_pos] / 2**0.5

        # Левая граница
        left = np.where(TR[:peak_pos] <= half_height)[0]
        if len(left) > 0 and (peak_pos - left[-1]) >= 1:
            TR_left = TR[left[-1]:peak_pos+1]
            freqs_left = freqs[left[-1]:peak_pos+1]
            if len(TR_left) >= 2 and len(freqs_left) >= 2:
                f1 = np.interp(half_height, TR_left[::-1], freqs_left[::-1])
            else:
                f1 = freqs[left[-1]]
        else:
            f1 = freqs[0]

        # Правая граница
        right = np.where(TR[peak_pos:] <= half_height)[0]
        if len(right) > 0:
            right_end = peak_pos + right[0] + 1
            TR_right = TR[peak_pos:right_end]
            freqs_right = freqs[peak_pos:right_end]
            if len(TR_right) >= 2 and len(freqs_right) >= 2:
                f2 = np.interp(half_height, TR_right, freqs_right)
            else:
                f2 = freqs[right_end-1]
        else:
            f2 = freqs[-1]

        return f1, f2

    except Exception as e:
        print(f"[find_res_width2] Ошибка: {e}")
        return -1, -1

